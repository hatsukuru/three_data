from docx import Document
import requests
import scrapy
import time
import datetime
import re
from selenium.webdriver.common.by import By
from selenium.webdriver import chromium
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver import ChromeOptions
from selenium import webdriver
from bs4 import BeautifulSoup
from datetime import datetime

class ThreeDatas8Spider(scrapy.Spider):
    name = 'three_datas8'
    allowed_domains = ['www.google.co.uk']
    start_urls = ['http://www.google.co.uk/']
    uu = 'https://time.com/search/?q={}&page=%d'

    google = 'https://translate.google.com.hk/?sl=en&tl=zh-CN&text={}&op=translate'

    ues = ['www.cnn.com','reuters.com','www.bbc.com', 'www.theguardian.com', 'apnews.com']

    source1 = ['CNN','路透社','BBC', '卫报', '美联社']
    app1 = ['TikTok', 'Alipay', 'WeChat']
    APP1 = ['抖音','支付宝','微信']

    def parse(self, response):  # scrapy crawl three_datas8
        # print('ass')
        # print(urls)
        options = ChromeOptions()
        options.add_experimental_option('excludeSwitches', ['enable - logging'])
        self.driver = webdriver.Chrome(options=options)
        for i in range(0, 1):
            # print('as')
            content_app = self.app1[i]
            content_app1 = self.app1[i]
            for ii in range(0, 1):
                # print(content_app)
                content_ues = self.ues[ii]
                content_source1 = self.source1[ii]
                # alls = content_app + '+site%3A' + content_ues
                url = 'https://www.google.co.uk/search?as_q={}'.format(content_app)


                date_now__ = datetime.now()
                date_now = str(date_now__.date())
                month = date_now__.month
                day = date_now__.day
                print(day)
                print(day-1)
                # print(date_now)

#https://www.google.co.uk/search?as_q=tiktok&as_epq=&as_oq=&as_eq=&as_nlo=&as_nhi=&lr=&cr=&as_qdr=d&as_sitesearch=reuters.com&as_occt=title&safe=active&as_filetype=&tbs=
#https://www.google.co.uk/search?as_q=tiktok&
                ui = f'&as_epq=&as_oq=&as_eq=&as_nlo=&as_nhi=&lr=&cr=&as_qdr=d&as_sitesearch={content_ues}&as_occt=title&safe=active&as_filetype=&tbs='
                urls = url + ui
                print(urls)
                self.driver.get(urls)

                original_window = self.driver.current_window_handle
                doc = Document()
                p = doc.add_paragraph()
                print('dsasd')


                time.sleep(5)
                try:
                    for iiii in range(1, 5):
                        path = '//*[@id="rso"]/div/div[{}]'.format(iiii)
                        path_one = '/div/div/div[1]/div/a'
                        print(path+path_one)
                        ele_source = self.driver.find_element(By.XPATH, path+path_one)
                        data_href = ele_source.get_attribute("href")
                        print(data_href)

                        try:
                            self.driver.get(data_href)
                        except Exception:
                            self.driver.execute_script('window.stop()')
                        titles = self.driver.title
                        print(titles)
                        try: ####                           英文文章获取
                            content = self.driver.find_element(By.XPATH, '//article').get_attribute('outerHTML')
                            soup = BeautifulSoup(content, 'html.parser')  # 对html进行解析
                            content_english1 = soup.get_text()
                            con = re.sub(r'\n', '', content_english1)
                            content_english2 = re.sub(r'\t', '', con)
                            content_english = content_english2.replace('  ','')
                        except Exception:
                            content = self.driver.find_element(By.XPATH, '//main').get_attribute('outerHTML')
                            soup = BeautifulSoup(content, 'html.parser')  # 对html进行解析
                            content_english1 = soup.get_text()
                            con = re.sub(r'\n', '', content_english1)
                            content_english2 = re.sub(r'\t', '', con)
                            content_english = content_english2.replace('  ','')
                        print(content_english)

                        urlsss = 'https://translate.google.com/?sl=en&tl=zh-CN'

                        self.driver.execute_script(f'window.open("{urlsss}");')
                        time.sleep(6)
                        self.driver.switch_to.window(self.driver.window_handles[-1])
                        print('ASas')

                        path_input = '//textarea[@aria-label="原文"]'
                        print('xzzZX')
                        self.driver.find_element(By.XPATH, path_input).send_keys(content_english)
                        print('asdasdasdasdasda')
                        time.sleep(30)
                        pathss = '//*[@id="yDmH0d"]/c-wiz/div/div[2]/c-wiz/div[2]/c-wiz/div[1]/div[2]/div[3]/c-wiz[2]/div/div[8]/div/div[1]/span[1]'
                                 #//*[@id="yDmH0d"]/c-wiz/div/div[2]/c-wiz/div[2]/c-wiz/div[1]/div[2]/div[3]/c-wiz[2]/div/div[8]/div/div[1]/span[1]
                        content1 = self.driver.find_element(By.XPATH, pathss).get_attribute('outerHTML')
                        soup1 = BeautifulSoup(content1, 'html.parser')  # 对html进行解析
                        content_chinese = soup1.get_text()
                        print(content_chinese)
                        print(content_source1)
                        print(data_href)
                        print(content_app)
                        print('ddddddddddddddddddddddddd')

                        source_title = '\n来源:' + content_source1
                        print(source_title)
                        source_href = '\n原文链接:' + data_href
                        print(source_href)
                        source_app = '\nAPP:' + content_app
                        print(source_app)
                        source_english = '\n英文内容:' + content_english
                        print("source_english")
                        source_chinese = '\n中文内容:' + content_chinese
                        print("source_chinese")

                        p.add_run(source_title)
                        print(source_title)
                        doc.add_page_break()
                        p.add_run(source_href)
                        print(source_href)
                        doc.add_page_break()
                        p.add_run(source_app)
                        print(source_app)
                        doc.add_page_break()
                        p.add_run(source_english)
                        doc.add_page_break()
                        p.add_run(source_chinese)
                        print("asdcxkjllk")
                        doc.save(iiii + content_source1 + content_app + date_now + '.doc')


                        self.driver.close()
                        self.driver.switch_to.window(original_window)

                        self.driver.back()



                except:
                    pass
        self.driver.quit()
